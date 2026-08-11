from django.shortcuts import render, redirect
from django.contrib.auth.models import User
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, JsonResponse
from django.conf import settings

from urllib.parse import quote

from .models import UserProfile, Activity

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

from docx import Document

from ai_tools.gemini import generate_text

from gtts import gTTS

import requests
import time
import json
import os
import uuid


# =========================================================
# ACTIVITY / HISTORY HELPER
# =========================================================

def log_activity(request, activity_type, tool_name, description=""):
    """
    Saves activity so it can appear on History page.

    This helper is written to work with different Activity model
    field names where possible.
    """

    try:
        user = request.user

        # Basic data
        activity_data = {}

        # User field
        try:
            Activity._meta.get_field("user")
            activity_data["user"] = user
        except Exception:
            pass

        # Activity type
        try:
            Activity._meta.get_field("activity_type")
            activity_data["activity_type"] = activity_type
        except Exception:
            pass

        # Try common possible field names for tool name
        possible_name_fields = [
            "tool_name",
            "name",
            "title",
            "activity_name",
        ]

        for field_name in possible_name_fields:
            try:
                Activity._meta.get_field(field_name)
                activity_data[field_name] = tool_name
                break
            except Exception:
                continue

        # Try common possible description fields
        possible_description_fields = [
            "description",
            "details",
            "message",
            "content",
        ]

        for field_name in possible_description_fields:
            try:
                Activity._meta.get_field(field_name)
                activity_data[field_name] = description
                break
            except Exception:
                continue

        Activity.objects.create(**activity_data)

    except Exception as e:
        # History problem should never stop the actual AI tool
        print("Activity logging error:", e)


# =========================================================
# LANDING PAGE
# =========================================================

def landing(request):
    return render(request, "index.html")


# =========================================================
# REGISTER
# =========================================================

def register_page(request):

    if request.method == "POST":

        fullname = request.POST.get("fullname", "").strip()
        email = request.POST.get("email", "").strip()
        password = request.POST.get("password", "")
        confirm = request.POST.get("confirmPassword", "")

        if not fullname or not email or not password:
            messages.error(request, "Please fill all required fields.")
            return redirect("register")

        if password != confirm:
            messages.error(request, "Passwords do not match.")
            return redirect("register")

        if User.objects.filter(username=email).exists():
            messages.error(request, "Email already exists.")
            return redirect("register")

        user = User.objects.create_user(
            username=email,
            email=email,
            password=password,
            first_name=fullname
        )

        user.save()

        UserProfile.objects.get_or_create(
            user=user
        )

        messages.success(
            request,
            "Registration Successful!"
        )

        return redirect("login")

    return render(request, "register.html")


def login_page(request):

    if request.method == "POST":

        email = request.POST.get("email", "").strip()
        password = request.POST.get("password", "")

        user = authenticate(
            username=email,
            password=password
        )

        if user is not None:

            login(request, user)

            return redirect("dashboard")

        messages.error(
            request,
            "Invalid Email or Password"
        )

    return render(request, "login.html")


def logout_user(request):

    logout(request)

    return redirect("login")


@login_required(login_url="login")
def dashboard(request):

    activities = Activity.objects.filter(
        user=request.user
    )

    document_count = activities.filter(
        activity_type="document"
    ).count()

    tts_count = activities.filter(
        activity_type="tts"
    ).count()

    image_count = activities.filter(
        activity_type="image"
    ).count()

    video_count = activities.filter(
        activity_type="video"
    ).count()

    recent_activities = activities.order_by(
        "-created_at"
    )[:5]

    return render(
        request,
        "dashboard.html",
        {
            "document_count": document_count,
            "tts_count": tts_count,
            "image_count": image_count,
            "video_count": video_count,
            "recent_activities": recent_activities,
        }
    )


@login_required(login_url="login")
def history(request):

    activities = Activity.objects.filter(
        user=request.user
    ).order_by(
        "-created_at"
    )

    return render(
        request,
        "history.html",
        {
            "activities": activities,
        }
    )


@login_required(login_url="login")
def settings_page(request):

    return render(
        request,
        "settings.html"
    )


@login_required(login_url="login")
def profile(request):

    profile, created = UserProfile.objects.get_or_create(
        user=request.user
    )

    if request.method == "POST":

        fullname = request.POST.get(
            "fullname",
            ""
        ).strip()

        if fullname:

            request.user.first_name = fullname

            request.user.save()

        profile.phone = request.POST.get(
            "phone",
            ""
        ).strip()

        profile.location = request.POST.get(
            "location",
            ""
        ).strip()

        uploaded_image = request.FILES.get(
            "profile_image"
        )

        if uploaded_image:

            profile.profile_image = uploaded_image

        profile.save()

        messages.success(
            request,
            "Profile updated successfully!"
        )

        return redirect("profile")

    activities = Activity.objects.filter(
        user=request.user
    ).order_by(
        "-created_at"
    )

    return render(
        request,
        "profile.html",
        {
            "profile": profile,
            "activities": activities,
        }
    )


@login_required(login_url="login")
def ai_tools(request):

    return render(
        request,
        "ai-tools.html"
    )


@login_required(login_url="login")
def text_generator(request):

    ai_response = ""

    if request.method == "POST":

        prompt = request.POST.get(
            "prompt",
            ""
        ).strip()

        if prompt:

            try:

                ai_response = generate_text(prompt)

                log_activity(
                    request,
                    "document",
                    "Text Generator",
                    "Generated AI text from a custom prompt."
                )

            except Exception as e:

                ai_response = (
                    f"Unable to generate text: {str(e)}"
                )

        else:

            ai_response = "Please enter a prompt."

    return render(
        request,
        "text-generator.html",
        {
            "response": ai_response
        }
    )


@login_required(login_url="login")
def text_to_speech(request):

    if request.method == "POST":

        text = request.POST.get(
            "text",
            ""
        ).strip()

        if not text:

            return JsonResponse(
                {
                    "error": "No text provided"
                },
                status=400
            )

        try:

            filename = f"{uuid.uuid4()}.mp3"

            folder = os.path.join(
                settings.MEDIA_ROOT,
                "tts"
            )

            os.makedirs(
                folder,
                exist_ok=True
            )

            filepath = os.path.join(
                folder,
                filename
            )

            tts = gTTS(
                text=text,
                lang="en"
            )

            tts.save(filepath)

            audio_url = (
                settings.MEDIA_URL
                + "tts/"
                + filename
            )

            log_activity(
                request,
                "tts",
                "Text To Speech",
                "Converted text into speech."
            )

            return JsonResponse(
                {
                    "audio_url": audio_url
                }
            )

        except Exception as e:

            return JsonResponse(
                {
                    "error": str(e)
                },
                status=500
            )

    return render(
        request,
        "text-to-speech.html"
    )


@login_required(login_url="login")
def speech_to_text(request):

    if request.method == "POST":

        log_activity(
            request,
            "document",
            "Speech To Text",
            "Used Speech To Text tool."
        )

    return render(
        request,
        "speech-to-text.html"
    )


@login_required(login_url="login")
def image_generator(request):

    image_url = None

    if request.method == "POST":

        prompt = request.POST.get(
            "prompt",
            ""
        ).strip()

        if prompt:

            encoded_prompt = quote(
                prompt
            )

            image_url = (
                f"https://image.pollinations.ai/prompt/"
                f"{encoded_prompt}"
            )

            log_activity(
                request,
                "image",
                "AI Image Generator",
                f"Generated image using prompt: {prompt[:100]}"
            )

    return render(
        request,
        "image-generator.html",
        {
            "image_url": image_url
        }
    )


@login_required(login_url="login")
def resume_builder(request):

    ai_response = ""

    ats_score = None

    score_breakdown = {}

    strengths = []

    weaknesses = []

    missing_keywords = []

    suggestions = []

    if request.method == "POST":

        fullname = request.POST.get("fullname", "")
        email = request.POST.get("email", "")
        phone = request.POST.get("phone", "")
        role = request.POST.get("role", "")
        summary = request.POST.get("summary", "")
        education = request.POST.get("education", "")
        skills = request.POST.get("skills", "")
        projects = request.POST.get("projects", "")
        experience = request.POST.get("experience", "")

    

        resume_prompt = f"""
You are an expert Resume Writer.

Create a professional ATS-friendly resume.

Full Name: {fullname}
Email: {email}
Phone: {phone}
Target Role: {role}

Professional Summary:
{summary}

Education:
{education}

Skills:
{skills}

Projects:
{projects}

Experience:
{experience}

Generate the resume using these sections:

1. Professional Summary
2. Education
3. Skills
4. Projects
5. Experience

Use professional formatting.
Use strong action verbs.
Make it ATS Friendly.
"""

        try:

            ai_response = generate_text(
                resume_prompt
            )

            request.session["resume"] = ai_response

            

            ats_prompt = f"""
You are an ATS Resume Analyzer.

Analyze the following resume professionally.

Return ONLY valid JSON.

{{
    "ats_score": 0,

    "score_breakdown": {{
        "resume_format": 0,
        "keywords": 0,
        "skills": 0,
        "experience": 0,
        "education": 0,
        "readability": 0
    }},

    "strengths": [],

    "weaknesses": [],

    "missing_keywords": [],

    "suggestions": []
}}

Rules:

ATS Score = out of 100

Resume Format = out of 10
Keywords = out of 20
Skills = out of 20
Experience = out of 20
Education = out of 10
Readability = out of 20

Return ONLY JSON.

Resume:

{ai_response}
"""

            try:

                ats_result = generate_text(
                    ats_prompt
                )

                ats_result = (
                    ats_result
                    .replace("```json", "")
                    .replace("```", "")
                    .strip()
                )

                ats_data = json.loads(
                    ats_result
                )

                ats_score = ats_data.get(
                    "ats_score",
                    0
                )

                score_breakdown = ats_data.get(
                    "score_breakdown",
                    {}
                )

                strengths = ats_data.get(
                    "strengths",
                    []
                )

                weaknesses = ats_data.get(
                    "weaknesses",
                    []
                )

                missing_keywords = ats_data.get(
                    "missing_keywords",
                    []
                )

                suggestions = ats_data.get(
                    "suggestions",
                    []
                )

            except Exception as e:

                ats_score = 0

                score_breakdown = {
                    "resume_format": 0,
                    "keywords": 0,
                    "skills": 0,
                    "experience": 0,
                    "education": 0,
                    "readability": 0
                }

                suggestions = [
                    "ATS Analysis Failed.",
                    str(e)
                ]

            log_activity(
                request,
                "document",
                "Resume Builder",
                "Generated and analyzed an ATS-friendly resume."
            )

        except Exception as e:

            ai_response = (
                f"Resume generation failed: {str(e)}"
            )

    return render(
        request,
        "resume-builder.html",
        {
            "response": ai_response,
            "ats_score": ats_score,
            "score_breakdown": score_breakdown,
            "strengths": strengths,
            "weaknesses": weaknesses,
            "missing_keywords": missing_keywords,
            "suggestions": suggestions,
        }
    )


@login_required(login_url="login")
def download_resume_pdf(request):

    resume = request.session.get(
        "resume",
        ""
    )

    if not resume:

        return HttpResponse(
            "No resume available."
        )

    response = HttpResponse(
        content_type="application/pdf"
    )

    response[
        "Content-Disposition"
    ] = 'attachment; filename="AI_Resume.pdf"'

    p = canvas.Canvas(
        response,
        pagesize=letter
    )

    text = p.beginText(
        40,
        750
    )

    text.setFont(
        "Helvetica",
        11
    )

    for line in resume.split("\n"):

        text.textLine(
            line
        )

    p.drawText(
        text
    )

    p.save()

    return response


@login_required(login_url="login")
def download_resume_docx(request):

    resume = request.session.get(
        "resume",
        ""
    )

    if not resume:

        return HttpResponse(
            "No resume available."
        )

    document = Document()

    document.add_heading(
        "AI Generated Resume",
        level=1
    )

    for line in resume.split("\n"):

        document.add_paragraph(
            line
        )

    response = HttpResponse(
        content_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        )
    )

    response[
        "Content-Disposition"
    ] = 'attachment; filename="AI_Resume.docx"'

    document.save(
        response
    )

    return response


@login_required(login_url="login")
def translator(request):

    ai_response = ""

    if request.method == "POST":

        text = request.POST.get(
            "text",
            ""
        ).strip()

        language = request.POST.get(
            "language",
            ""
        ).strip()

        if text and language:

            prompt = f"""
Translate the following text into {language}.

Only return the translated text.

Text:
{text}
"""

            try:

                ai_response = generate_text(
                    prompt
                )

                log_activity(
                    request,
                    "document",
                    "Translator",
                    f"Translated text into {language}."
                )

            except Exception as e:

                ai_response = (
                    f"Translation failed: {str(e)}"
                )

        else:

            ai_response = (
                "Please enter text and select a language."
            )

    return render(
        request,
        "translator.html",
        {
            "response": ai_response
        }
    )


@login_required(login_url="login")
def chatbot(request):

    user_message = ""

    ai_response = ""

    if request.method == "POST":

        user_message = request.POST.get(
            "message",
            ""
        ).strip()

        if user_message:

            prompt = f"""
You are SmartAI Assistant.

Reply in a friendly, helpful and professional way.

User:
{user_message}
"""

            try:

                ai_response = generate_text(
                    prompt
                )

                log_activity(
                    request,
                    "document",
                    "AI Chatbot",
                    "Used SmartAI Chatbot."
                )

            except Exception as e:

                ai_response = (
                    f"Chatbot error: {str(e)}"
                )

    return render(
        request,
        "chatbot.html",
        {
            "user_message": user_message,
            "ai_response": ai_response
        }
    )


@login_required(login_url="login")
def ai_video_generator(request):

    video_url = None

    status = ""

    if request.method == "POST":

        prompt = request.POST.get(
            "prompt",
            ""
        ).strip()

        if not prompt:

            status = (
                "Please enter a video prompt."
            )

        else:

            try:


                magic_hour_url = (
                    "https://api.magichour.ai/v1/text-to-video"
                )

                headers = {
                    "Accept": "application/json",
                    "Authorization": (
                        f"Bearer "
                        f"{settings.MAGIC_HOUR_API_KEY}"
                    ),
                    "Content-Type": "application/json",
                }

                data = {
                    "name": "SmartAI Generated Video",
                    "end_seconds": 1,
                    "orientation": "landscape",
                    "resolution": "480p",
                    "style": {
                        "prompt": prompt
                    }
                }

                print("\n==============================")
                print("MAGIC HOUR REQUEST")
                print("PROMPT:", prompt)
                print("==============================")

                response = requests.post(
                    magic_hour_url,
                    headers=headers,
                    json=data,
                    timeout=60
                )

                print(
                    "STATUS:",
                    response.status_code
                )

                print(
                    "RESPONSE:",
                    response.text
                )


                if response.status_code not in [
                    200,
                    201
                ]:

                    try:

                        error_data = response.json()

                    except Exception:

                        error_data = response.text

                    status = (
                        f"❌ Magic Hour API Error "
                        f"({response.status_code}): "
                        f"{error_data}"
                    )

                else:

                    result = response.json()

                    project_id = result.get(
                        "id"
                    )

                    print(
                        "PROJECT ID:",
                        project_id
                    )

                    if not project_id:

                        status = (
                            "❌ Magic Hour did not "
                            "return a project ID."
                        )

                    else:

                        status = (
                            "🎬 Video generation started..."
                        )


                        project_url = (
                            "https://api.magichour.ai/"
                            f"v1/video-projects/{project_id}"
                        )

                        for attempt in range(60):

                            time.sleep(5)

                            status_response = requests.get(
                                project_url,
                                headers={
                                    "Accept": "application/json",
                                    "Authorization": (
                                        f"Bearer "
                                        f"{settings.MAGIC_HOUR_API_KEY}"
                                    )
                                },
                                timeout=30
                            )

                            print(
                                "VIDEO STATUS:",
                                status_response.status_code
                            )

                            print(
                                "VIDEO RESPONSE:",
                                status_response.text
                            )

                            if status_response.status_code != 200:

                                status = (
                                    "❌ Unable to check "
                                    "video status: "
                                    f"{status_response.text}"
                                )

                                break

                            video_data = (
                                status_response.json()
                            )

                            current_status = (
                                video_data.get(
                                    "status"
                                )
                            )

                            print(
                                "CURRENT STATUS:",
                                current_status
                            )

                            if current_status == "complete":

                                downloads = (
                                    video_data.get(
                                        "downloads",
                                        []
                                    )
                                )

                                if downloads:

                                    video_url = (
                                        downloads[0].get(
                                            "url"
                                        )
                                    )

                                    if video_url:

                                        status = (
                                            "🎉 Video Generated "
                                            "Successfully!"
                                        )

                                        log_activity(
                                            request,
                                            "video",
                                            "AI Video Generator",
                                            "Generated an AI video successfully."
                                        )

                                    else:

                                        status = (
                                            "❌ Video completed "
                                            "but video URL "
                                            "is empty."
                                        )

                                else:

                                    status = (
                                        "❌ Video completed "
                                        "but download URL "
                                        "was not found."
                                    )

                                break


                            elif current_status == "error":

                                error_data = (
                                    video_data.get(
                                        "error"
                                    )
                                )

                                status = (
                                    "❌ Video generation "
                                    "failed: "
                                    f"{error_data}"
                                )

                                break


                            elif current_status == "canceled":

                                status = (
                                    "❌ Video generation "
                                    "was canceled."
                                )

                                break


                            else:

                                status = (
                                    f"🎬 Video is "
                                    f"{current_status}..."
                                )

                        else:

                            status = (
                                "⏳ Video is taking longer "
                                "than expected. Please try "
                                "again later."
                            )


            except requests.exceptions.RequestException as e:

                status = (
                    "❌ Magic Hour connection error: "
                    f"{str(e)}"
                )

            except Exception as e:

                status = (
                    "❌ Video generation error: "
                    f"{str(e)}"
                )

    return render(
        request,
        "ai-video-generator.html",
        {
            "video_url": video_url,
            "status": status,
        }
    )